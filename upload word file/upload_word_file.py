from PyPDF2 import PdfReader
from docx import Document
import os

def upload_file(file_path):
    if not os.path.exists(file_path):
        return 'File not found'

    if file_path.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        text = 'Unsupported file format'

    return text

def extract_text_from_pdf(pdf_file):
    with open(pdf_file, 'rb') as f:
        reader = PdfReader(f)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def main():from flask import Flask, render_template, request
from PyPDF2 import PdfReader
from flask import Flask, render_template, request

from docx import Document
import os

app = Flask(__name__)

@app.route('/')
def upload_file():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_and_extract_text():
    if request.method == 'POST':
        file = request.files['file']
        if file.filename == '':
            return 'No selected file'
        if file:
            file_path = file.filename
            file.save(file_path)
            text = process_file(file_path)
            os.remove(file_path)
            return text

def process_file(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        return 'Unsupported file format'

def extract_text_from_pdf(pdf_file):
    with open(pdf_file, 'rb') as f:
        reader = PdfReader(f)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def main():
    app.run(debug=True)

if __name__ == "__main__":
    main()

    file_path = input("Enter the file path: ")
    result = upload_file(file_path)
    print("Text extracted from file:\n", result)

if __name__ == "__main__":
    main()
